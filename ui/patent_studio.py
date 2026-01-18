"""
专利智能工作室 - Patent Studio

一站式专利交底书处理和专利文件生成平台
- 支持文件上传（Word、TXT、PDF）
- 智能内容润色
- 实时预览和编辑
- 一键生成专利申请文件
"""

import streamlit as st
import asyncio
import os
import sys
import json
import re
from datetime import datetime
from typing import Optional, Dict, Any, List
import traceback

# 可选依赖 - 文件处理
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    docx = None

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    PyPDF2 = None

# 添加父目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 页面配置
st.set_page_config(
    page_title="专利智能工作室 | Patent Studio",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded",
)


def init_session_state():
    """初始化会话状态"""
    defaults = {
        "current_page": "home",
        "disclosure_data": {},
        "uploaded_content": "",
        "generated_patent": "",
        "patent_sections": {},
        "generation_progress": {},
        "polished_sections": {},
        "processing_status": None,
        "step": 1,
        "optimized_document": "",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def parse_uploaded_file(uploaded_file) -> str:
    """解析上传的文件"""
    if uploaded_file is None:
        return ""

    file_type = uploaded_file.name.split(".")[-1].lower()

    try:
        if file_type == "txt":
            content = uploaded_file.read().decode("utf-8")
            return clean_extracted_text(content)

        elif file_type in ["doc", "docx"]:
            if not HAS_DOCX:
                st.error("缺少 python-docx 模块，请安装: pip install python-docx")
                return ""
            from io import BytesIO

            file_content = uploaded_file.getvalue()
            doc = docx.Document(BytesIO(file_content))

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)

            return clean_extracted_text("\n".join(text_parts))

        elif file_type == "pdf":
            if not HAS_PDF:
                st.error("缺少 PyPDF2 模块，请安装: pip install PyPDF2")
                return ""
            from io import BytesIO

            file_content = uploaded_file.getvalue()
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))

            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text.strip())

            return clean_extracted_text("\n".join(text_parts))

        else:
            st.error(f"不支持的文件格式: {file_type}")
            return ""

    except Exception as e:
        st.error(f"读取文件失败: {str(e)}")
        return ""


def clean_extracted_text(text: str) -> str:
    """清理提取的文本，保持格式统一"""
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

    return text.strip()


async def extract_disclosure_fields_with_llm(text: str) -> Dict[str, str]:
    """使用 LLM 智能提取交底书字段"""
    system_prompt = """你是一位专业的专利代理人，擅长从技术交底书中提取和整理信息。

请仔细阅读用户提供的技术交底书内容，准确提取以下信息：

1. **发明名称** (title): 技术的正式名称
2. **技术领域** (technical_field): 技术所属的领域
3. **申请人** (applicant_name): 申请人或申请单位名称
4. **发明人** (inventors): 发明人姓名列表
5. **背景技术** (background_description): 现有技术状况和存在的问题
6. **要解决的技术问题** (technical_problems): 本发明要解决的具体技术问题
7. **技术方案** (technical_solution): 本发明的技术方案详细描述
8. **有益效果** (beneficial_effects): 与现有技术相比的有益效果
9. **具体实施例** (embodiments): 具体的实施方式和实施例
10. **附图说明** (figure_descriptions): 附图的简要说明

注意事项：
- 如果某个字段在交底书中没有明确提及，请根据上下文合理推断
- 保持提取内容的完整性和准确性，不要过度简化
- 对于技术方案和背景技术，尽可能保留详细描述
- 只返回 JSON 格式的结果，不要添加任何解释或说明

返回格式（必须是有效的 JSON）：
{
  "title": "...",
  "technical_field": "...",
  "applicant_name": "...",
  "inventors": "...",
  "background_description": "...",
  "technical_problems": "...",
  "technical_solution": "...",
  "beneficial_effects": "...",
  "embodiments": "...",
  "figure_descriptions": "..."
}"""

    max_length = 12000
    if len(text) > max_length:
        text = text[:max_length] + "\n...(内容过长，已截断)"

    prompt = f"""请从以下技术交底书内容中提取各字段信息：

【技术交底书内容】
{text}

请仔细阅读并提取所有相关信息，返回完整的 JSON 格式结果。"""

    try:
        response = await call_deepseek_api(prompt, system_prompt)

        import json

        response = response.strip()
        if response.startswith("```"):
            response = re.sub(r"^```(?:json)?\s*", "", response)
            response = re.sub(r"\s*```$", "", response)

        json_match = re.search(r"\{[\s\S]*\}", response)
        if json_match:
            response = json_match.group(0)

        parsed = json.loads(response)

        default_fields = {
            "title": "",
            "technical_field": "",
            "applicant_name": "",
            "inventors": "",
            "background_description": "",
            "technical_problems": "",
            "technical_solution": "",
            "beneficial_effects": "",
            "embodiments": "",
            "figure_descriptions": "",
        }

        default_fields.update(parsed)
        return default_fields

    except json.JSONDecodeError as e:
        st.error(f"LLM 返回的内容格式有误: {str(e)}")
        return extract_disclosure_fields_fallback(text)
    except Exception as e:
        st.error(f"LLM 提取失败: {str(e)}")
        return extract_disclosure_fields_fallback(text)


def extract_disclosure_fields_fallback(text: str) -> Dict[str, str]:
    """降级方法：使用正则表达式提取交底书字段"""
    fields = {
        "title": "",
        "technical_field": "",
        "background_description": "",
        "technical_problems": "",
        "technical_solution": "",
        "beneficial_effects": "",
        "embodiments": "",
        "inventors": "",
        "applicant_name": "",
        "figure_descriptions": "",
    }

    patterns = {
        "title": [r"发明名称[：:]\s*(.+?)(?:\n|$)", r"名称[：:]\s*(.+?)(?:\n|$)"],
        "technical_field": [
            r"技术领域[：:]\s*(.+?)(?:\n\n|\n[一二三四五六七八九十]|$)",
            r"所属领域[：:]\s*(.+?)(?:\n|$)",
        ],
        "background_description": [
            r"背景技术[：:]\s*(.+?)(?:\n\n[一二三四五六七八九十]|要解决|技术问题|$)"
        ],
        "technical_problems": [
            r"(?:要解决的)?技术问题[：:]\s*(.+?)(?:\n\n|技术方案|$)",
            r"解决的问题[：:]\s*(.+?)(?:\n|$)",
        ],
        "technical_solution": [
            r"技术方案[：:]\s*(.+?)(?:\n\n[一二三四五六七八九十]|有益效果|$)"
        ],
        "beneficial_effects": [
            r"有益效果[：:]\s*(.+?)(?:\n\n[一二三四五六七八九十]|实施例|附图|$)"
        ],
        "embodiments": [
            r"(?:具体)?实施例[：:]\s*(.+?)(?:\n\n[一二三四五六七八九十]|附图|$)"
        ],
        "inventors": [r"发明人[：:]\s*(.+?)(?:\n|$)"],
        "applicant_name": [r"申请人[：:]\s*(.+?)(?:\n|$)"],
        "figure_descriptions": [
            r"附图说明[：:]\s*(.+?)(?:\n\n[一二三四五六七八九十]|具体实施方式|$)",
            r"图[：:]\s*(.+?)(?:\n|$)",
        ],
    }

    for field, pattern_list in patterns.items():
        for pattern in pattern_list:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                fields[field] = match.group(1).strip()
                break

    return fields


async def call_deepseek_api(prompt: str, system_prompt: str = None) -> str:
    """调用 DeepSeek API"""
    try:
        from openai import OpenAI
        import streamlit as st

        api_key = st.secrets.get("DEEPSEEK_API_KEY", os.environ.get("DEEPSEEK_API_KEY"))

        client = OpenAI(
            api_key=api_key,
            base_url="https://api.deepseek.com/v1"
        )

        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})

        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            temperature=0.7,
            max_tokens=4096,
        )

        return response.choices[0].message.content

    except Exception as e:
        return f"API 调用失败: {str(e)}"


async def polish_content(content: str, field_name: str) -> str:
    """润色指定内容"""
    field_prompts = {
        "title": "请润色以下专利发明名称，使其更加专业、准确、符合专利命名规范（不超过25个字）：",
        "technical_field": "请润色以下技术领域描述，使其更加准确、专业：",
        "background_description": "请润色以下背景技术描述，使其更加详细、专业，并突出现有技术的问题和不足：",
        "technical_problems": "请润色以下技术问题描述，使其更加清晰、具体：",
        "technical_solution": "请润色以下技术方案描述，使其更加详细、专业、逻辑清晰：",
        "beneficial_effects": "请润色以下有益效果描述，使其更加具体、可量化：",
        "embodiments": "请润色以下实施例描述，使其更加详细、可操作：",
    }

    system_prompt = """你是一位专业的专利代理人，擅长撰写高质量的专利申请文件。
请根据要求润色内容，保持原意的同时提升专业性和规范性。
只返回润色后的内容，不要添加任何解释或说明。"""

    prompt = field_prompts.get(field_name, "请润色以下内容：") + f"\n\n{content}"

    return await call_deepseek_api(prompt, system_prompt)


async def generate_patent_section(
    section_name: str,
    disclosure_data: Dict[str, str],
    context: Dict[str, str] = None
) -> str:
    """生成专利的某个章节"""

    section_prompts = {
        "title": {
            "system": "你是专利代理人，擅长撰写符合中国专利规范的发明名称。",
            "prompt": f"""根据以下技术交底书，撰写一个规范、准确的发明名称。
要求：
1. 简明扼要，全面反映发明的技术主题和类型
2. 不超过25个字
3. 不使用人名、地名、商标或商业性宣传用语
4. 避免使用"新型"、"改进"等词

技术交底书：
{disclosure_data.get('title', '')}
{disclosure_data.get('technical_solution', '')[:200]}

请只返回发明名称，不要其他内容。"""
        },

        "technical_field": {
            "system": "你是专利代理人，擅长撰写技术领域章节。",
            "prompt": f"""根据以下技术交底书，撰写技术领域章节。
要求：
1. 明确指出发明所属的技术领域
2. 可以是某个较大的技术领域，也可以是其中的一个分支
3. 通常用"本发明涉及..."或"本发明属于..."开头
4. 长度适中，通常1-2段

技术交底书：
发明名称：{disclosure_data.get('title', '')}
技术领域：{disclosure_data.get('technical_field', '')}

请只返回技术领域章节内容。"""
        },

        "background": {
            "system": "你是专利代理人，擅长撰写背景技术章节。",
            "prompt": f"""根据以下技术交底书，撰写背景技术章节。
要求：
1. 描述与本发明相关的现有技术状况
2. 指出现有技术存在的问题和不足
3. 说明解决这些问题对技术进步的意义
4. 客观引用参考文献或现有技术
5. 长度适中，通常2-4段

技术交底书：
背景技术：{disclosure_data.get('background_description', '')}
技术问题：{disclosure_data.get('technical_problems', '')}

请只返回背景技术章节内容。"""
        },

        "invention_content": {
            "system": "你是专利代理人，擅长撰写发明内容章节。",
            "prompt": f"""根据以下技术交底书，撰写发明内容章节。
要求：
1. 要解决的技术问题：清晰、具体
2. 技术方案：详细、完整，使本领域技术人员能够实现
3. 有益效果：具体、可验证，与现有技术对比

技术交底书：
技术问题：{disclosure_data.get('technical_problems', '')}
技术方案：{disclosure_data.get('technical_solution', '')}
有益效果：{disclosure_data.get('beneficial_effects', '')}

请按以下格式返回：

【要解决的技术问题】
...

【技术方案】
...

【有益效果】
..."""
        },

        "claims": {
            "system": "你是专利代理人，擅长撰写权利要求书。权利要求书是专利的核心，必须精确、完整。",
            "prompt": f"""根据以下技术交底书，撰写权利要求书。
要求：
1. 必须包含1项独立权利要求和9项从属权利要求（共10项）
2. 独立权利要求应包含发明的全部必要技术特征
3. 从属权利要求应进一步限定独立权利要求或引用在前权利要求
4. 使用"根据权利要求X所述的..."格式
5. 技术特征描述清晰、准确
6. 保护范围合理，既不过宽也不过窄

技术交底书：
发明名称：{disclosure_data.get('title', '')}
技术方案：{disclosure_data.get('technical_solution', '')}
有益效果：{disclosure_data.get('beneficial_effects', '')}
实施例：{disclosure_data.get('embodiments', '')}

请按以下格式返回：

1. 一种[发明名称]，其特征在于，包括：...

2. 根据权利要求1所述的[发明名称]，其特征在于：...

3. 根据权利要求1或2所述的[发明名称]，其特征在于：...

（共10项权利要求）"""
        },

        "embodiments": {
            "system": "你是专利代理人，擅长撰写具体实施方式章节。",
            "prompt": f"""根据以下技术交底书，撰写具体实施方式章节。
要求：
1. 详细描述至少1个具体实施例
2. 实施例应充分公开发明的技术方案
3. 包含具体的参数、步骤、结构等细节
4. 使本领域技术人员能够实现
5. 可以包含多个不同的实施例

技术交底书：
技术方案：{disclosure_data.get('technical_solution', '')}
实施例：{disclosure_data.get('embodiments', '')}
附图说明：{disclosure_data.get('figure_descriptions', '')}

请返回详细的具体实施方式章节内容。"""
        },

        "abstract": {
            "system": "你是专利代理人，擅长撰写摘要。",
            "prompt": f"""根据以下技术交底书，撰写摘要。
要求：
1. 简明扼要地介绍发明的技术领域、技术方案和有益效果
2. 字数控制在300字以内
3. 不包含商业性宣传用语
4. 客观反映发明的技术要点

技术交底书：
发明名称：{disclosure_data.get('title', '')}
技术方案：{disclosure_data.get('technical_solution', '')}
有益效果：{disclosure_data.get('beneficial_effects', '')}

请只返回摘要内容。"""
        },

        "figure_description": {
            "system": "你是专利代理人，擅长撰写附图说明。",
            "prompt": f"""根据以下技术交底书，撰写附图说明章节。
要求：
1. 列出所有附图的图名和简要说明
2. 按照附图的顺序编号（图1、图2、...）
3. 说明每个附图展示的内容
4. 简洁明了，一句话说明即可

技术交底书：
附图说明：{disclosure_data.get('figure_descriptions', '')}
技术方案：{disclosure_data.get('technical_solution', '')[:500]}

请按以下格式返回：

图1是...的示意图/流程图/结构图；
图2是...的示意图；
...
（如果没有附图，请返回"本申请不包含附图。"）"""
        }
    }

    config = section_prompts.get(section_name)
    if not config:
        return ""

    return await call_deepseek_api(config["prompt"], config["system"])


async def optimize_patent_document(sections: Dict[str, str]) -> str:
    """总览并优化整个专利文件"""

    sections_text = "\n\n".join([
        f"【{name}】\n{content}"
        for name, content in sections.items()
    ])

    system_prompt = """你是一位资深专利代理人，擅长审核和优化专利申请文件。
请仔细阅读以下专利申请文件的所有章节，指出存在的问题并提出修改建议。

重点关注：
1. 各章节之间的一致性（术语、编号等）
2. 权利要求书与说明书的一致性
3. 技术方案描述的完整性和清晰性
4. 有益效果是否与技术方案对应
5. 是否存在遗漏或需要补充的内容

请提供：
1. 整体评价
2. 具体问题列表（按严重程度排序）
3. 修改建议
4. 优化后的完整专利文件"""

    prompt = f"""请审核并优化以下专利申请文件：

{sections_text}

请先分析问题，然后提供优化后的完整专利文件。"""

    return await call_deepseek_api(prompt, system_prompt)


def render_home_page():
    """渲染首页"""
    st.title("⚡ 专利智能工作室")
    st.markdown("**AI 驱动的一站式专利文件生成平台 | Powered by DeepSeek**")

    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.info("📤 **智能上传**\n\n支持 Word、TXT、PDF 等多种格式，自动识别和提取交底书内容")

    with col2:
        st.success("✨ **AI 润色**\n\n对每个章节进行智能润色，提升专利文件的专业性和规范性")

    with col3:
        st.warning("📄 **一键生成**\n\n自动生成符合规范的完整专利申请文件，包括权利要求书")

    st.markdown("---")
    st.markdown("### 🚀 快速开始")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("📤 上传交底书", use_container_width=True, type="primary"):
            st.session_state.current_page = "upload"
            st.rerun()

    with col2:
        if st.button("📝 在线填写", use_container_width=True):
            st.session_state.current_page = "fill"
            st.rerun()


def render_upload_page():
    """渲染上传页面"""
    st.header("📤 上传交底书文件")

    if st.button("← 返回首页"):
        st.session_state.current_page = "home"
        st.rerun()

    st.markdown("---")

    uploaded_file = st.file_uploader(
        "选择交底书文件",
        type=["txt", "doc", "docx", "pdf"],
        help="支持 TXT、Word（.doc/.docx）、PDF 格式"
    )

    if uploaded_file:
        with st.spinner("正在解析文件..."):
            content = parse_uploaded_file(uploaded_file)

            if content:
                st.success(f"✅ 文件解析成功！共 {len(content)} 个字符")

                st.info("🤖 正在使用 AI 智能识别交底书内容...")

                with st.spinner("正在使用 DeepSeek AI 智能提取内容..."):
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    extracted = loop.run_until_complete(extract_disclosure_fields_with_llm(content))
                    loop.close()

                st.session_state.disclosure_data = extracted

                with st.expander("📄 查看原文", expanded=False):
                    st.text_area("文件内容", content, height=300, disabled=True)

                st.markdown("### 📋 AI 识别结果")
                st.success("以下内容由 AI 智能识别提取，您可以编辑和润色后生成专利文件")

                tab1, tab2 = st.tabs(["📝 基本信息", "📄 技术内容"])

                with tab1:
                    col1, col2 = st.columns(2)
                    with col1:
                        title = st.text_input("发明名称", value=extracted.get("title", ""))
                        applicant = st.text_input("申请人", value=extracted.get("applicant_name", ""))
                    with col2:
                        inventors = st.text_input("发明人", value=extracted.get("inventors", ""))
                        tech_field = st.text_input("技术领域", value=extracted.get("technical_field", ""))

                    st.session_state.disclosure_data.update({
                        "title": title,
                        "applicant_name": applicant,
                        "inventors": inventors,
                        "technical_field": tech_field,
                    })

                with tab2:
                    st.markdown("#### 背景技术")
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        bg = st.text_area(
                            "背景技术描述",
                            value=extracted.get("background_description", ""),
                            height=150,
                            key="bg_text"
                        )
                    with col2:
                        if st.button("✨ 润色", key="polish_bg"):
                            with st.spinner("润色中..."):
                                loop = asyncio.new_event_loop()
                                asyncio.set_event_loop(loop)
                                polished = loop.run_until_complete(polish_content(bg, "background_description"))
                                loop.close()
                                st.session_state.disclosure_data["background_description"] = polished
                                st.rerun()

                    st.session_state.disclosure_data["background_description"] = bg

                    st.markdown("#### 要解决的技术问题")
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        problems = st.text_area(
                            "技术问题",
                            value=extracted.get("technical_problems", ""),
                            height=100,
                            key="problems_text"
                        )
                    with col2:
                        if st.button("✨ 润色", key="polish_problems"):
                            with st.spinner("润色中..."):
                                loop = asyncio.new_event_loop()
                                asyncio.set_event_loop(loop)
                                polished = loop.run_until_complete(polish_content(problems, "technical_problems"))
                                loop.close()
                                st.session_state.disclosure_data["technical_problems"] = polished
                                st.rerun()

                    st.session_state.disclosure_data["technical_problems"] = problems

                    st.markdown("#### 技术方案")
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        solution = st.text_area(
                            "技术方案",
                            value=extracted.get("technical_solution", ""),
                            height=200,
                            key="solution_text"
                        )
                    with col2:
                        if st.button("✨ 润色", key="polish_solution"):
                            with st.spinner("润色中..."):
                                loop = asyncio.new_event_loop()
                                asyncio.set_event_loop(loop)
                                polished = loop.run_until_complete(polish_content(solution, "technical_solution"))
                                loop.close()
                                st.session_state.disclosure_data["technical_solution"] = polished
                                st.rerun()

                    st.session_state.disclosure_data["technical_solution"] = solution

                    st.markdown("#### 有益效果")
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        effects = st.text_area(
                            "有益效果",
                            value=extracted.get("beneficial_effects", ""),
                            height=100,
                            key="effects_text"
                        )
                    with col2:
                        if st.button("✨ 润色", key="polish_effects"):
                            with st.spinner("润色中..."):
                                loop = asyncio.new_event_loop()
                                asyncio.set_event_loop(loop)
                                polished = loop.run_until_complete(polish_content(effects, "beneficial_effects"))
                                loop.close()
                                st.session_state.disclosure_data["beneficial_effects"] = polished
                                st.rerun()

                    st.session_state.disclosure_data["beneficial_effects"] = effects

                    st.markdown("#### 具体实施例")
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        embodiments = st.text_area(
                            "具体实施例",
                            value=extracted.get("embodiments", ""),
                            height=150,
                            key="embodiments_text",
                            placeholder="描述具体实施方式..."
                        )
                    with col2:
                        if st.button("✨ 润色", key="polish_embodiments"):
                            with st.spinner("润色中..."):
                                loop = asyncio.new_event_loop()
                                asyncio.set_event_loop(loop)
                                polished = loop.run_until_complete(polish_content(embodiments, "embodiments"))
                                loop.close()
                                st.session_state.disclosure_data["embodiments"] = polished
                                st.rerun()

                    st.session_state.disclosure_data["embodiments"] = embodiments

                    st.markdown("#### 附图说明")
                    figures = st.text_area(
                        "附图说明",
                        value=extracted.get("figure_descriptions", ""),
                        height=80,
                        key="figures_text",
                        placeholder="例如：图1为系统架构图；图2为流程图..."
                    )
                    st.session_state.disclosure_data["figure_descriptions"] = figures

                st.markdown("---")
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button("🚀 生成专利申请文件", type="primary", use_container_width=True):
                        st.session_state.current_page = "generate"
                        st.rerun()


def render_fill_page():
    """渲染在线填写页面"""
    st.header("📝 在线填写交底书")

    if st.button("← 返回首页"):
        st.session_state.current_page = "home"
        st.rerun()

    st.markdown("---")

    step = st.session_state.step
    steps = ["基本信息", "技术内容", "确认生成"]

    cols = st.columns(len(steps) * 2 - 1)
    for i, step_name in enumerate(steps):
        with cols[i * 2]:
            if i + 1 < step:
                st.markdown(f"✅ **{step_name}**")
            elif i + 1 == step:
                st.markdown(f"🔵 **{step_name}**")
            else:
                st.markdown(f"⚪ {step_name}")
        if i < len(steps) - 1:
            with cols[i * 2 + 1]:
                st.markdown("—")

    st.markdown("---")

    data = st.session_state.disclosure_data

    if step == 1:
        st.markdown("#### 📌 基本信息")

        col1, col2 = st.columns(2)

        with col1:
            title = st.text_input(
                "发明名称 *",
                value=data.get("title", ""),
                placeholder="例如：一种基于深度学习的智能推荐方法"
            )

            patent_type = st.selectbox(
                "专利类型 *",
                options=["invention", "utility_model", "design"],
                format_func=lambda x: {"invention": "发明专利", "utility_model": "实用新型", "design": "外观设计"}[x]
            )

            technical_field = st.text_input(
                "技术领域 *",
                value=data.get("technical_field", ""),
                placeholder="例如：人工智能、机器学习"
            )

        with col2:
            applicant_name = st.text_input(
                "申请人名称 *",
                value=data.get("applicant_name", ""),
                placeholder="公司名称或个人姓名"
            )

            applicant_address = st.text_input(
                "申请人地址",
                value=data.get("applicant_address", ""),
                placeholder="详细地址"
            )

            inventors = st.text_input(
                "发明人 *",
                value=data.get("inventors", ""),
                placeholder="多个发明人用逗号分隔"
            )

        st.session_state.disclosure_data.update({
            "title": title,
            "patent_type": patent_type,
            "technical_field": technical_field,
            "applicant_name": applicant_name,
            "applicant_address": applicant_address,
            "inventors": inventors,
        })

        is_valid = all([title, technical_field, applicant_name, inventors])

        col1, col2, col3 = st.columns([1, 1, 1])
        with col3:
            if st.button("下一步 →", type="primary", use_container_width=True, disabled=not is_valid):
                st.session_state.step = 2
                st.rerun()

        if not is_valid:
            st.warning("请填写所有必填项（标记 * 的字段）")

    elif step == 2:
        st.markdown("#### 📄 技术内容")

        st.markdown("##### 背景技术 *")
        col1, col2 = st.columns([6, 1])
        with col1:
            background = st.text_area(
                "描述现有技术及存在的问题",
                value=data.get("background_description", ""),
                height=150,
                placeholder="请详细描述当前技术领域的发展状况和存在的问题..."
            )
        with col2:
            if st.button("✨ 润色", key="p_bg"):
                if background:
                    with st.spinner("润色中..."):
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        polished = loop.run_until_complete(polish_content(background, "background_description"))
                        loop.close()
                        st.session_state.disclosure_data["background_description"] = polished
                        st.rerun()
        st.session_state.disclosure_data["background_description"] = background

        st.markdown("##### 要解决的技术问题 *")
        col1, col2 = st.columns([6, 1])
        with col1:
            problems = st.text_area(
                "本发明要解决的具体技术问题",
                value=data.get("technical_problems", ""),
                height=100,
                placeholder="例如：如何提高系统的处理速度和准确率..."
            )
        with col2:
            if st.button("✨ 润色", key="p_prob"):
                if problems:
                    with st.spinner("润色中..."):
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        polished = loop.run_until_complete(polish_content(problems, "technical_problems"))
                        loop.close()
                        st.session_state.disclosure_data["technical_problems"] = polished
                        st.rerun()
        st.session_state.disclosure_data["technical_problems"] = problems

        st.markdown("##### 技术方案 *")
        col1, col2 = st.columns([6, 1])
        with col1:
            solution = st.text_area(
                "详细描述技术方案",
                value=data.get("technical_solution", ""),
                height=200,
                placeholder="请详细描述您的技术方案，包括主要步骤和实现方式..."
            )
        with col2:
            if st.button("✨ 润色", key="p_sol"):
                if solution:
                    with st.spinner("润色中..."):
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        polished = loop.run_until_complete(polish_content(solution, "technical_solution"))
                        loop.close()
                        st.session_state.disclosure_data["technical_solution"] = polished
                        st.rerun()
        st.session_state.disclosure_data["technical_solution"] = solution

        st.markdown("##### 有益效果 *")
        col1, col2 = st.columns([6, 1])
        with col1:
            effects = st.text_area(
                "与现有技术相比的有益效果",
                value=data.get("beneficial_effects", ""),
                height=100,
                placeholder="例如：提高效率30%、降低成本50%..."
            )
        with col2:
            if st.button("✨ 润色", key="p_eff"):
                if effects:
                    with st.spinner("润色中..."):
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        polished = loop.run_until_complete(polish_content(effects, "beneficial_effects"))
                        loop.close()
                        st.session_state.disclosure_data["beneficial_effects"] = polished
                        st.rerun()
        st.session_state.disclosure_data["beneficial_effects"] = effects

        st.markdown("##### 具体实施例")
        embodiments = st.text_area(
            "提供具体的实施例（可选）",
            value=data.get("embodiments", ""),
            height=150,
            placeholder="描述一个或多个具体的实施例..."
        )
        st.session_state.disclosure_data["embodiments"] = embodiments

        st.markdown("##### 附图说明")
        figures = st.text_area(
            "附图说明（可选）",
            value=data.get("figure_descriptions", ""),
            height=80,
            placeholder="图1为系统架构图；图2为流程图..."
        )
        st.session_state.disclosure_data["figure_descriptions"] = figures

        is_valid = all([background, problems, solution, effects])

        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("← 上一步", use_container_width=True):
                st.session_state.step = 1
                st.rerun()
        with col3:
            if st.button("下一步 →", type="primary", use_container_width=True, disabled=not is_valid):
                st.session_state.step = 3
                st.rerun()

        if not is_valid:
            st.warning("请填写所有必填项")

    elif step == 3:
        st.markdown("#### ✅ 确认信息")

        data = st.session_state.disclosure_data

        patent_type_map = {'invention': '发明专利', 'utility_model': '实用新型', 'design': '外观设计'}
        patent_type_name = patent_type_map.get(data.get('patent_type', ''), '')

        with st.expander("📋 交底书信息预览", expanded=True):
            st.markdown(f"""
            **发明名称**: {data.get('title', '')}

            **专利类型**: {patent_type_name}

            **技术领域**: {data.get('technical_field', '')}

            **申请人**: {data.get('applicant_name', '')}

            **发明人**: {data.get('inventors', '')}

            ---

            **背景技术**:
            {data.get('background_description', '')[:200]}...

            **技术问题**:
            {data.get('technical_problems', '')[:200]}...

            **技术方案**:
            {data.get('technical_solution', '')[:200]}...

            **有益效果**:
            {data.get('beneficial_effects', '')}
            """)

        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            if st.button("← 返回修改", use_container_width=True):
                st.session_state.step = 2
                st.rerun()
        with col2:
            if st.button("🚀 生成专利申请文件", type="primary", use_container_width=True):
                st.session_state.current_page = "generate"
                st.rerun()


def render_generate_page():
    """渲染生成页面 - 分板块生成专利"""
    st.header("📄 生成专利申请文件")

    col1, col2 = st.columns([1, 5])
    with col1:
        if st.button("← 返回"):
            st.session_state.current_page = "home"
            st.session_state.step = 1
            st.session_state.patent_sections = {}
            st.session_state.generation_progress = {}
            st.session_state.optimized_document = ""
            st.rerun()

    st.markdown("---")

    data = st.session_state.disclosure_data

    required_fields = ["title", "technical_field", "background_description", "technical_problems", "technical_solution", "beneficial_effects"]
    missing_fields = [f for f in required_fields if not data.get(f)]

    if missing_fields:
        st.error(f"缺少必要信息: {', '.join(missing_fields)}")
        st.warning("请返回完善交底书信息")
        return

    sections_config = {
        "title": {"name": "发明名称", "icon": "📌"},
        "technical_field": {"name": "技术领域", "icon": "🎯"},
        "background": {"name": "背景技术", "icon": "📚"},
        "invention_content": {"name": "发明内容", "icon": "💡"},
        "claims": {"name": "权利要求书（10条）", "icon": "⚖️"},
        "embodiments": {"name": "具体实施方式", "icon": "🔧"},
        "figure_description": {"name": "附图说明", "icon": "🖼️"},
        "abstract": {"name": "摘要", "icon": "📝"},
    }

    if not st.session_state.patent_sections:
        st.info("👇 点击下方按钮开始分板块生成专利文件")

        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚀 开始生成专利文件", type="primary", use_container_width=True):
                st.session_state.generation_progress = {}
                st.session_state.patent_sections = {}

                progress_bar = st.progress(0)
                status_text = st.empty()
                current_section = st.empty()

                total_sections = len(sections_config)
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

                for i, (section_key, section_info) in enumerate(sections_config.items()):
                    status_text.text(f"正在生成：{section_info['name']}")
                    current_section.markdown(f"**当前步骤：** {section_info['icon']} {section_info['name']}")

                    try:
                        result = loop.run_until_complete(
                            generate_patent_section(section_key, data)
                        )
                        st.session_state.patent_sections[section_key] = result
                        st.session_state.generation_progress[section_key] = "completed"
                    except Exception as e:
                        st.session_state.patent_sections[section_key] = f"生成失败: {str(e)}"
                        st.session_state.generation_progress[section_key] = "failed"

                    progress = (i + 1) / total_sections
                    progress_bar.progress(progress)

                loop.close()

                progress_bar.empty()
                status_text.empty()
                current_section.empty()

                st.success("✅ 所有章节生成完成！")
                st.rerun()

    else:
        st.success("✅ 专利文件章节已生成完成")

        st.markdown("### 📋 章节概览")

        for section_key, section_info in sections_config.items():
            content = st.session_state.patent_sections.get(section_key, "")

            if "生成失败" in content:
                status = "❌"
            elif content:
                status = "✅"
            else:
                status = "⏳"

            with st.expander(f"{status} {section_info['icon']} {section_info['name']}", expanded=False):
                if content and "生成失败" not in content:
                    st.text_area(
                        f"{section_info['name']}内容",
                        value=content,
                        height=200,
                        key=f"section_{section_key}"
                    )

                    if st.button(f"🔄 重新生成{section_info['name']}", key=f"regen_{section_key}"):
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        result = loop.run_until_complete(
                            generate_patent_section(section_key, data)
                        )
                        loop.close()
                        st.session_state.patent_sections[section_key] = result
                        st.rerun()
                else:
                    st.error(content)
                    if st.button(f"🔄 重试{section_info['name']}", key=f"retry_{section_key}"):
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        result = loop.run_until_complete(
                            generate_patent_section(section_key, data)
                        )
                        loop.close()
                        st.session_state.patent_sections[section_key] = result
                        st.rerun()

        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🔍 总览并优化整篇专利文件", type="primary", use_container_width=True):
                with st.spinner("正在分析并优化专利文件..."):
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    result = loop.run_until_complete(
                        optimize_patent_document(st.session_state.patent_sections)
                    )
                    loop.close()
                    st.session_state.optimized_document = result
                st.success("✅ 优化完成！")
                st.rerun()

        if st.session_state.optimized_document:
            st.markdown("---")
            st.markdown("### 🎯 优化后的专利文件")

            tab1, tab2 = st.tabs(["📄 优化结果", "📥 下载"])

            with tab1:
                st.text(st.session_state.optimized_document)

            with tab2:
                filename = f"专利申请文件_{data.get('title', '未命名')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

                col1, col2 = st.columns(2)

                with col1:
                    st.download_button(
                        label="📄 下载 TXT 格式",
                        data=st.session_state.optimized_document,
                        file_name=f"{filename}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )

                with col2:
                    export_data = {
                        "disclosure": data,
                        "sections": st.session_state.patent_sections,
                        "optimized_document": st.session_state.optimized_document,
                        "generated_at": datetime.now().isoformat(),
                        "model": "DeepSeek",
                    }
                    st.download_button(
                        label="📦 下载 JSON 格式",
                        data=json.dumps(export_data, ensure_ascii=False, indent=2),
                        file_name=f"{filename}.json",
                        mime="application/json",
                        use_container_width=True,
                    )

        st.markdown("---")
        col1, col2, col3 = st.columns([1, 1, 1])

        with col1:
            if st.button("📝 返回修改", use_container_width=True):
                st.session_state.patent_sections = {}
                st.session_state.optimized_document = ""
                st.session_state.current_page = "fill"
                st.session_state.step = 2
                st.rerun()

        with col2:
            if st.button("🔄 重新生成所有", use_container_width=True):
                st.session_state.patent_sections = {}
                st.session_state.optimized_document = ""
                st.rerun()

        with col3:
            if st.button("🏠 返回首页", use_container_width=True):
                st.session_state.disclosure_data = {}
                st.session_state.patent_sections = {}
                st.session_state.optimized_document = ""
                st.session_state.step = 1
                st.session_state.current_page = "home"
                st.rerun()


def render_sidebar():
    """渲染侧边栏"""
    with st.sidebar:
        st.markdown("### ⚡ 专利智能工作室")
        st.markdown("---")

        st.markdown("#### 📋 导航")

        if st.button("🏠 首页", use_container_width=True):
            st.session_state.current_page = "home"
            st.rerun()

        if st.button("📤 上传交底书", use_container_width=True):
            st.session_state.current_page = "upload"
            st.rerun()

        if st.button("📝 在线填写", use_container_width=True):
            st.session_state.current_page = "fill"
            st.session_state.step = 1
            st.rerun()

        st.markdown("---")

        st.markdown("#### 🤖 AI 模型")
        st.markdown("""
        <span style='color: green;'>● </span> **在线**

        **DeepSeek Chat**

        <small style='color: gray;'>高性能中文大模型</small>
        """, unsafe_allow_html=True)

        st.markdown("---")

        with st.expander("❓ 使用说明"):
            st.markdown("""
            **功能说明：**

            1. **上传交底书**
               - 支持 Word、TXT、PDF
               - 自动识别和提取内容

            2. **在线填写**
               - 分步骤引导填写
               - 实时保存进度

            3. **AI 润色**
               - 点击 ✨ 按钮润色内容
               - 提升专业性和规范性

            4. **生成专利文件**
               - 自动生成完整文件
               - 包含权利要求书
            """)

        st.markdown("---")
        st.markdown(
            '<p style="text-align: center; color: gray; font-size: 0.75rem;">'
            'Powered by DeepSeek AI<br>v2.0.0'
            '</p>',
            unsafe_allow_html=True
        )


def main():
    """主函数"""
    init_session_state()
    render_sidebar()

    page = st.session_state.current_page

    if page == "home":
        render_home_page()
    elif page == "upload":
        render_upload_page()
    elif page == "fill":
        render_fill_page()
    elif page == "generate":
        render_generate_page()


if __name__ == "__main__":
    main()
